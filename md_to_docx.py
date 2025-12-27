import docx
from docx.shared import Pt, Inches
import re
import os

def parse_markdown(md_path, docx_path):
    print(f"Converting {md_path} to {docx_path}...")
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    doc = docx.Document()
    
    # Simple parser state
    in_table = False
    table_lines = []
    
    for line in lines:
        line = line.strip()
        
        # 1. Handle Tables (Very basic pipe table support)
        if line.startswith('|'):
            in_table = True
            table_lines.append(line)
            continue
        else:
            if in_table:
                # Process the accumulated table
                process_table(doc, table_lines)
                table_lines = []
                in_table = False
        
        if not line:
            continue
            
        # 2. Handle Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[2:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[2:], level=3)
        
        # 3. Handle Lists
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.', line):
            # Ordered list
            text = re.sub(r'^\d+\.\s*', '', line)
            doc.add_paragraph(text, style='List Number')
            
        # 4. Normal Text
        else:
            doc.add_paragraph(line)
            
    # Flush last table if exists
    if in_table and table_lines:
        process_table(doc, table_lines)
        
    doc.save(docx_path)
    print(f"Saved {docx_path}")

def process_table(doc, lines):
    # Determine columns from first line
    # | Col1 | Col2 | -> ['', ' Col1 ', ' Col2 ', '']
    header = [c.strip() for c in lines[0].split('|') if c.strip() != '']
    if not header:
        return

    # Skip separator line if it looks like |---|---|
    data_lines = []
    for l in lines[1:]:
        if set(l).issubset({'|', '-', ' ', ':'}):
            continue
        data_lines.append(l)

    table = doc.add_table(rows=1, cols=len(header))
    table.style = 'Table Grid'
    
    # Set Header
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = h.strip()
        # Bold header
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True
            
    # Add Rows
    for l in data_lines:
        row_cells = table.add_row().cells
        # Split but handle potential missing or extra pipes? 
        # For this specific report, it's clean.
        cols = [c.strip() for c in l.strip().split('|')][1:-1] # Remove first empty and last empty from split('|...|')
        
        for i in range(min(len(cols), len(header))):
            row_cells[i].text = cols[i]

if __name__ == "__main__":
    MD_FILE = "C:/Users/mizal/.gemini/antigravity/brain/13ac5187-4bb2-406a-9545-e0fbed5bb711/expert_discussion_report.md"
    DOCX_FILE = "C:/Users/mizal/.gemini/antigravity/brain/13ac5187-4bb2-406a-9545-e0fbed5bb711/expert_discussion_report.docx"
    
    parse_markdown(MD_FILE, DOCX_FILE)
